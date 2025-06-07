"""
VedaVMS DOCX Parser
Extracts text from Baraha-encoded DOCX files from vedavms.in
"""

import sys
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import re
import logging

try:
    from docx import Document
except ImportError:
    print("python-docx not found. Install with: pip install python-docx")
    sys.exit(1)

logger = logging.getLogger(__name__)

class VedavmsParser:
    """Parser for VedaVMS DOCX files with Baraha encoding."""
    
    def __init__(self):
        self.source_name = "vedavms"
        self.encoding = "baraha"
        
    def parse_file(self, file_path: str) -> Dict:
        """
        Parse a VedaVMS DOCX file and extract structured text.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dict with extracted text and metadata
        """
        file_path = Path(file_path)
        
        try:
            doc = Document(file_path)
            
            # Extract metadata from filename
            metadata = self._parse_filename(file_path.name)
            
            # Extract text content
            text_content = self._extract_text(doc)
            
            # Parse structure from content
            structure = self._parse_structure(text_content, metadata)
            
            result = {
                'source_file': str(file_path),
                'source_name': self.source_name,
                'encoding': self.encoding,
                'metadata': metadata,
                'raw_text': text_content,
                'structure': structure,
                'extraction_method': 'python-docx',
                'file_size': file_path.stat().st_size,
                'file_modified': file_path.stat().st_mtime
            }
            
            logger.info(f"Successfully parsed {file_path.name}: {len(text_content)} characters")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            raise
    
    def _parse_filename(self, filename: str) -> Dict:
        """Extract metadata from VedaVMS filename patterns."""
        metadata = {
            'veda': 'yajurveda',
            'sakha': 'taittiriya',
            'original_filename': filename
        }
        
        # Parse different filename patterns
        if filename.startswith('TS '):
            # Taittiriya Samhita patterns
            metadata['text_type'] = 'samhita'
            
            if 'Pada paatam' in filename:
                # Individual adhyaya files: "TS X.Y Baraha Pada paatam.docx"
                match = re.search(r'TS (\d+)\.(\d+)', filename)
                if match:
                    metadata['kanda'] = int(match.group(1))
                    metadata['prapathaka'] = int(match.group(2))
                    metadata['level'] = 'prapathaka'
            else:
                # Kanda-level files: "TS X Baraha.docx"
                match = re.search(r'TS (\d+)', filename)
                if match:
                    metadata['kanda'] = int(match.group(1))
                    metadata['level'] = 'kanda'
                    
        elif filename.startswith('TB '):
            # Taittiriya Brahmana: "TB X.Y-X.Z Baraha.docx"
            metadata['text_type'] = 'brahmana'
            match = re.search(r'TB (\d+)\.(\d+)-(\d+)\.(\d+)', filename)
            if match:
                metadata['kanda_start'] = int(match.group(1))
                metadata['prapathaka_start'] = int(match.group(2))
                metadata['kanda_end'] = int(match.group(3))
                metadata['prapathaka_end'] = int(match.group(4))
                metadata['level'] = 'prapathaka_range'
                
        elif filename.startswith('TA '):
            # Taittiriya Aranyaka: "TA X-Y Baraha.docx"
            metadata['text_type'] = 'aranyaka'
            match = re.search(r'TA (\d+)-(\d+)', filename)
            if match:
                metadata['prapathaka_start'] = int(match.group(1))
                metadata['prapathaka_end'] = int(match.group(2))
                metadata['level'] = 'prapathaka_range'
        
        return metadata
    
    def _extract_text(self, doc: Document) -> str:
        """Extract text from DOCX document, preserving structure."""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        return '\n'.join(text_parts)
    
    def _parse_structure(self, text: str, metadata: Dict) -> Dict:
        """Parse hierarchical structure from text content."""
        structure = {
            'veda': metadata['veda'],
            'sakha': metadata['sakha'],
            'text_type': metadata.get('text_type'),
            'sections': []
        }
        
        # Split text into sections based on common patterns
        # This is a basic implementation - can be enhanced based on actual file analysis
        lines = text.split('\n')
        current_section = None
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Look for section markers (numbers, traditional markers, etc.)
            section_match = re.search(r'^(\d+)[\.:]', line)
            if section_match:
                if current_section:
                    structure['sections'].append(current_section)
                
                current_section = {
                    'section_number': int(section_match.group(1)),
                    'line_start': line_num,
                    'content': [line]
                }
            elif current_section:
                current_section['content'].append(line)
            else:
                # Content before first section
                if 'preamble' not in structure:
                    structure['preamble'] = []
                structure['preamble'].append(line)
        
        # Add the last section
        if current_section:
            structure['sections'].append(current_section)
        
        return structure
    
    def get_file_list(self, base_path: str) -> List[str]:
        """Get list of VedaVMS DOCX files to process."""
        base_path = Path(base_path)
        docx_files = []
        
        # Find all DOCX files in VedaVMS directories
        for vedavms_dir in base_path.glob('**/vedavms/*/source'):
            for docx_file in vedavms_dir.glob('*.docx'):
                # Skip temporary files
                if not docx_file.name.startswith('~'):
                    docx_files.append(str(docx_file))
        
        return sorted(docx_files)
    
    def validate_file(self, file_path: str) -> bool:
        """Validate that file can be processed."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return False
            
        if file_path.suffix.lower() != '.docx':
            logger.error(f"Not a DOCX file: {file_path}")
            return False
            
        if file_path.name.startswith('~'):
            logger.warning(f"Temporary file, skipping: {file_path}")
            return False
            
        try:
            # Try to open the document
            doc = Document(file_path)
            return True
        except Exception as e:
            logger.error(f"Cannot open DOCX file {file_path}: {e}")
            return False


if __name__ == "__main__":
    # Test the parser
    parser = VedavmsParser()
    
    # Find test files
    base_path = "/Users/skmnktl/Projects/udapaana/data/vedic_texts"
    files = parser.get_file_list(base_path)
    
    print(f"Found {len(files)} VedaVMS DOCX files:")
    for file_path in files[:5]:  # Show first 5
        print(f"  {Path(file_path).name}")
        
        if parser.validate_file(file_path):
            try:
                result = parser.parse_file(file_path)
                print(f"    ✓ Parsed: {len(result['raw_text'])} chars, {result['metadata']}")
            except Exception as e:
                print(f"    ✗ Error: {e}")
        else:
            print(f"    ✗ Validation failed")
        print()