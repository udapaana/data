#!/usr/bin/env python3
"""
Convert VedaVMS DOCX files to Udapaana JSON format.
Handles Baraha encoding and preserves accent marks for vidyut-lipi processing.
"""

import json
import sys
import os
from pathlib import Path
from typing import Dict, List
import logging
import re

try:
    from docx import Document
except ImportError:
    print("python-docx not found. Install with: pip install python-docx")
    sys.exit(1)

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class VedavmsToJsonConverter:
    """Convert VedaVMS DOCX files to Udapaana JSON format."""
    
    def __init__(self):
        self.source_name = "vedavms"
        self.encoding = "BarahaSouth"  # Baraha encoding with special accent marks
        
        # Baraha accent patterns to preserve
        self.baraha_accent_patterns = [
            '^',    # udātta (high tone)
            '.',    # anudātta (low tone) 
            '~',    # svarita (falling tone)
            '-',    # pluta
            '#',    # section markers
            'q', 'w', 'x', 'z',  # Special Baraha characters
            'Q', 'W', 'X', 'Z',  # Capital variants
            '¤',    # Special marker
            'ˆ',    # Accent marker
        ]
    
    def convert_directory(self, base_path: str):
        """Convert all DOCX files in vedavms directories."""
        base_path = Path(base_path)
        docx_files = []
        
        # Find all DOCX files
        vedavms_path = base_path / "yajurveda" / "taittiriya" / "vedavms"
        if not vedavms_path.exists():
            logger.error(f"VedaVMS path not found: {vedavms_path}")
            return
        
        for text_type in ['samhita', 'brahmana', 'aranyaka']:
            source_dir = vedavms_path / text_type / "source"
            if source_dir.exists():
                for docx_file in source_dir.glob("*.docx"):
                    if not docx_file.name.startswith('~'):
                        docx_files.append(docx_file)
        
        logger.info(f"Found {len(docx_files)} DOCX files to convert")
        
        # Convert each file
        for docx_file in docx_files:
            try:
                self.convert_file(docx_file)
            except Exception as e:
                logger.error(f"Failed to convert {docx_file}: {e}")
    
    def convert_file(self, docx_path: Path):
        """Convert a single DOCX file to JSON."""
        logger.info(f"Converting {docx_path.name}...")
        
        # Parse the DOCX file
        doc = Document(docx_path)
        
        # Extract metadata from filename
        metadata = self._parse_filename(docx_path.name)
        
        # Extract all text content
        texts = self._extract_texts(doc, metadata)
        
        if not texts:
            logger.warning(f"No texts extracted from {docx_path.name}")
            return
        
        # Create Udapaana format JSON
        udapaana_data = {
            "texts": texts,
            "metadata": {
                "veda": metadata['veda'],
                "sakha": metadata['sakha'],
                "text_type": metadata['text_type'],
                "source": self.source_name,
                "source_encoding": self.encoding,
                "original_file": docx_path.name,
                "has_accents": True,  # Baraha files have accent marks
                "accent_patterns": self.baraha_accent_patterns,
                "notes": f"Converted from DOCX with Baraha encoding. Contains special accent marks: {', '.join(self.baraha_accent_patterns[:6])}"
            }
        }
        
        # Save JSON file
        json_path = docx_path.with_suffix('.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(udapaana_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"✓ Saved {json_path.name} with {len(texts)} verses")
    
    def _parse_filename(self, filename: str) -> Dict:
        """Extract metadata from filename."""
        metadata = {
            'veda': 'yajurveda',
            'sakha': 'taittiriya',
            'source': self.source_name
        }
        
        # Determine text type and location
        if filename.startswith('TS '):
            metadata['text_type'] = 'samhita'
            match = re.search(r'TS (\d+)(?:\.(\d+))?', filename)
            if match:
                metadata['kanda'] = int(match.group(1))
                if match.group(2):
                    metadata['prapathaka'] = int(match.group(2))
                    
        elif filename.startswith('TB '):
            metadata['text_type'] = 'brahmana'
            match = re.search(r'TB (\d+)\.(\d+)', filename)
            if match:
                metadata['kanda'] = int(match.group(1))
                metadata['prapathaka'] = int(match.group(2))
                
        elif filename.startswith('TA '):
            metadata['text_type'] = 'aranyaka'
            match = re.search(r'TA (\d+)', filename)
            if match:
                metadata['prapathaka'] = int(match.group(1))
        
        return metadata
    
    def _extract_texts(self, doc: Document, metadata: Dict) -> List[Dict]:
        """Extract individual verses/texts from document."""
        texts = []
        verse_counter = 1
        
        # Extract all paragraphs
        all_text = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                all_text.append(text)
        
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        all_text.append(text)
        
        # Split into verses based on patterns
        current_verse = []
        for line in all_text:
            # Check if this is a verse separator (numbers, double pipes, etc.)
            if self._is_verse_separator(line):
                if current_verse:
                    # Save previous verse
                    verse_text = ' '.join(current_verse)
                    if self._is_valid_verse(verse_text):
                        texts.append({
                            "vaakya_pk": verse_counter,
                            "vaakya_sk": f"{metadata.get('kanda', 0)}.{metadata.get('prapathaka', 0)}.{verse_counter}",
                            "vaakya_text": verse_text,
                            "location": [
                                metadata.get('kanda', 0),
                                metadata.get('prapathaka', 0),
                                verse_counter
                            ]
                        })
                        verse_counter += 1
                    current_verse = []
                
                # Start new verse with this line if it contains content
                cleaned_line = self._clean_verse_separator(line)
                if cleaned_line:
                    current_verse = [cleaned_line]
            else:
                # Continue current verse
                current_verse.append(line)
        
        # Don't forget the last verse
        if current_verse:
            verse_text = ' '.join(current_verse)
            if self._is_valid_verse(verse_text):
                texts.append({
                    "vaakya_pk": verse_counter,
                    "vaakya_sk": f"{metadata.get('kanda', 0)}.{metadata.get('prapathaka', 0)}.{verse_counter}",
                    "vaakya_text": verse_text,
                    "location": [
                        metadata.get('kanda', 0),
                        metadata.get('prapathaka', 0),
                        verse_counter
                    ]
                })
        
        return texts
    
    def _is_verse_separator(self, line: str) -> bool:
        """Check if line is a verse separator."""
        # Common verse ending patterns
        if '।।' in line or '||' in line:
            return True
        
        # Lines that are just numbers (verse numbers)
        if re.match(r'^\d+\.?\s*$', line):
            return True
        
        # Lines ending with verse numbers
        if re.search(r'।।\s*\d+\s*।।$', line):
            return True
        
        return False
    
    def _clean_verse_separator(self, line: str) -> str:
        """Extract content from verse separator line."""
        # Remove standalone numbers
        if re.match(r'^\d+\.?\s*$', line):
            return ''
        
        # Remove verse numbers at the end
        line = re.sub(r'।।\s*\d+\s*।।$', '।।', line)
        
        return line.strip()
    
    def _is_valid_verse(self, text: str) -> bool:
        """Check if text is a valid verse."""
        # Must have some content
        if len(text.strip()) < 5:
            return False
        
        # Should contain at least some Devanagari or Baraha characters
        # (not just punctuation and numbers)
        content_chars = re.sub(r'[।॥\|\.\s\d]+', '', text)
        if len(content_chars) < 3:
            return False
        
        return True
    
    def analyze_baraha_patterns(self, base_path: str):
        """Analyze Baraha patterns in DOCX files for better encoding detection."""
        base_path = Path(base_path)
        pattern_counts = {}
        
        vedavms_path = base_path / "yajurveda" / "taittiriya" / "vedavms"
        
        for text_type in ['samhita', 'brahmana', 'aranyaka']:
            source_dir = vedavms_path / text_type / "source"
            if source_dir.exists():
                for docx_file in source_dir.glob("*.docx"):
                    if not docx_file.name.startswith('~'):
                        try:
                            doc = Document(docx_file)
                            text = '\n'.join(p.text for p in doc.paragraphs)
                            
                            # Count special characters
                            for char in text:
                                if ord(char) > 127 or char in '^.~-#qwxzQWXZ¤ˆ':
                                    pattern_counts[char] = pattern_counts.get(char, 0) + 1
                        except Exception as e:
                            logger.error(f"Error analyzing {docx_file}: {e}")
        
        # Report findings
        logger.info("\nBaraha pattern analysis:")
        for char, count in sorted(pattern_counts.items(), key=lambda x: x[1], reverse=True)[:20]:
            logger.info(f"  '{char}' (U+{ord(char):04X}): {count:,} occurrences")

if __name__ == "__main__":
    converter = VedavmsToJsonConverter()
    
    # First analyze patterns
    logger.info("Analyzing Baraha patterns in DOCX files...")
    converter.analyze_baraha_patterns("data/vedic_texts")
    
    # Then convert files
    logger.info("\nConverting VedaVMS DOCX files to JSON...")
    converter.convert_directory("data/vedic_texts")